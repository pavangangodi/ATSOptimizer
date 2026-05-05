from __future__ import annotations

from io import BytesIO


def markdown_to_docx(markdown: str) -> bytes:
    from docx import Document

    document = Document()
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph()
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
